"""
ReportWriter birim testleri — Faz 5a (HTML / DOCX / PDF).

Gerçek LLM çağrısı yapılmaz; SolverState sabit bir fixture ile beslenir.
"""

from __future__ import annotations

from pathlib import Path

import pytest

from iesolver.report import ReportWriter, write_report

# ---------------------------------------------------------------------------
# Fixture
# ---------------------------------------------------------------------------

SAMPLE_STATE = {
    "explicit_goal": "EOQ: D=10000, S=50, H=2 — find optimal order quantity.",
    "executive_summary": "## Executive Summary\n\nOptimal order quantity is **707 units**. Annual cost savings estimated at $2,000.",
    "technical_output": "### Method\n\nEOQ formula: Q* = sqrt(2DS/H)\n\n- D=10000, S=50, H=2\n- Q* = 707.1 units\n- Cycle length = 25.8 days",
    "action_directives": "1. Place orders of **707 units** each cycle.\n2. Set reorder point at 500 units.\n3. Review demand forecast quarterly.",
    "sensitivity_results": "Shadow price on demand: 0.014. A 10% demand increase raises total cost by 0.5%.",
    "execution_result": "Q* = 707.10678...",
    "is_valid": True,
    "figures": [],
    "metrics": {
        "intake":  {"llm_calls": 1, "tokens_in": 120, "tokens_out": 30, "cost_usd": 0.0001, "latency_ms": 450},
        "report":  {"llm_calls": 1, "tokens_in": 800, "tokens_out": 250, "cost_usd": 0.0008, "latency_ms": 1200},
    },
}


# ---------------------------------------------------------------------------
# write_report convenience function
# ---------------------------------------------------------------------------

def test_write_report_html(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "report.html", format="html")  # type: ignore[arg-type]
    assert out.exists()
    content = out.read_text(encoding="utf-8")
    assert "<!DOCTYPE html>" in content
    assert "Executive Summary" in content
    assert "707" in content


def test_write_report_docx(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "report.docx", format="docx")  # type: ignore[arg-type]
    assert out.exists()
    assert out.stat().st_size > 1_000  # gerçek DOCX dosyası


def test_write_report_pdf(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "report.pdf", format="pdf")  # type: ignore[arg-type]
    assert out.exists()
    header = out.read_bytes()[:4]
    assert header == b"%PDF"


# ---------------------------------------------------------------------------
# ReportWriter class
# ---------------------------------------------------------------------------

def test_report_writer_creates_parent_dirs(tmp_path: Path) -> None:
    nested = tmp_path / "a" / "b" / "c" / "r.html"
    ReportWriter(SAMPLE_STATE).write(nested, format="html")  # type: ignore[arg-type]
    assert nested.exists()


def test_report_writer_invalid_format(tmp_path: Path) -> None:
    with pytest.raises(ValueError, match="Unsupported format"):
        ReportWriter(SAMPLE_STATE).write(tmp_path / "x.xyz", format="xyz")  # type: ignore[arg-type]


# ---------------------------------------------------------------------------
# HTML content checks
# ---------------------------------------------------------------------------

def test_html_contains_goal(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    assert "EOQ" in out.read_text(encoding="utf-8")


def test_html_contains_sensitivity_section(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    assert "Sensitivity Analysis" in out.read_text(encoding="utf-8")


def test_html_no_sensitivity_when_absent(tmp_path: Path) -> None:
    state = {**SAMPLE_STATE, "sensitivity_results": None}
    out = write_report(state, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    assert "Sensitivity Analysis" not in out.read_text(encoding="utf-8")


def test_html_metrics_table(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    html = out.read_text(encoding="utf-8")
    assert "Performance Metrics" in html
    assert "intake" in html


def test_html_embeds_figure(tmp_path: Path) -> None:
    # PNG dosyası oluştur (1x1 piksel)
    png_bytes = (
        b"\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01"
        b"\x00\x00\x00\x01\x08\x02\x00\x00\x00\x90wS\xde\x00\x00"
        b"\x00\x0cIDATx\x9cc\xf8\x0f\x00\x00\x01\x01\x00\x05\x18"
        b"\xd8N\x00\x00\x00\x00IEND\xaeB`\x82"
    )
    fig = tmp_path / "chart.png"
    fig.write_bytes(png_bytes)
    state = {**SAMPLE_STATE, "figures": [fig]}
    out = write_report(state, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    assert "data:image/png;base64," in out.read_text(encoding="utf-8")


def test_html_validated_badge(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "r.html", format="html")  # type: ignore[arg-type]
    assert "Validated" in out.read_text(encoding="utf-8")


# ---------------------------------------------------------------------------
# DOCX content checks
# ---------------------------------------------------------------------------

def test_docx_readable(tmp_path: Path) -> None:
    from docx import Document  # type: ignore[import]
    out = write_report(SAMPLE_STATE, tmp_path / "r.docx", format="docx")  # type: ignore[arg-type]
    doc = Document(str(out))
    full_text = " ".join(p.text for p in doc.paragraphs)
    assert "Executive Summary" in full_text or any(
        "Executive" in p.text for t in doc.tables for r in t.rows for p in r.cells[0].paragraphs
        # tables not checked here; heading paragraphs suffice
    )


def test_docx_has_metrics_table(tmp_path: Path) -> None:
    from docx import Document  # type: ignore[import]
    out = write_report(SAMPLE_STATE, tmp_path / "r.docx", format="docx")  # type: ignore[arg-type]
    doc = Document(str(out))
    assert len(doc.tables) >= 1


# ---------------------------------------------------------------------------
# PDF content checks
# ---------------------------------------------------------------------------

def test_pdf_is_valid_pdf(tmp_path: Path) -> None:
    out = write_report(SAMPLE_STATE, tmp_path / "r.pdf", format="pdf")  # type: ignore[arg-type]
    raw = out.read_bytes()
    assert raw[:4] == b"%PDF"
    assert b"%%EOF" in raw or b"%%EOF\n" in raw or b"%%EOF\r" in raw


def test_pdf_minimal_state(tmp_path: Path) -> None:
    """Empty state'te crash olmamalı."""
    out = write_report({}, tmp_path / "r.pdf", format="pdf")  # type: ignore[arg-type]
    assert out.read_bytes()[:4] == b"%PDF"
